#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read markdown content
md_file = '/root/opencode/k8s-resource-analyzer/docs/开发文档.md'
with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

# Create DOCX
docx_file = '/root/opencode/k8s-resource-analyzer/docs/K8s资源使用率分析工具_开发文档.docx'
doc = Document()

def clean_markdown(text):
    text = re.sub(r'#+\s+', '', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

def is_heading(line):
    return line.startswith('#')

def get_heading_level(line):
    count = 0
    for c in line:
        if c == '#':
            count += 1
        else:
            break
    return count

lines = content.split('\n')
i = 0
in_code_block = False
code_content = []

while i < len(lines):
    line = lines[i]
    
    if line.strip().startswith('```'):
        if in_code_block:
            code_text = '\n'.join(code_content)
            p = doc.add_paragraph(code_text)
            p.style = 'Quote'
            code_content = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    
    if in_code_block:
        code_content.append(line)
        i += 1
        continue
    
    line = line.strip()
    
    if not line:
        i += 1
        continue
    
    if is_heading(line):
        level = get_heading_level(line)
        text = clean_markdown(line)
        
        if level == 1:
            heading = doc.add_heading(text, level=0)
        elif level == 2:
            heading = doc.add_heading(text, level=1)
        else:
            heading = doc.add_heading(text, level=2)
    else:
        if '|' in line and '-' not in line:
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if len(parts) >= 2:
                table_data = [parts]
                i += 1
                while i < len(lines):
                    row_line = lines[i].strip()
                    if not row_line or '|' not in row_line:
                        break
                    row_parts = [p.strip() for p in row_line.split('|') if p.strip()]
                    if len(row_parts) == len(parts):
                        table_data.append(row_parts)
                    i += 1
                
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_data):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        row.cells[col_idx].text = cell_text
                
                doc.add_paragraph()
                continue
        else:
            line = clean_markdown(line)
            if line.startswith('• '):
                line = line.replace('• ', '')
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
    
    i += 1

doc.save(docx_file)
print(f"DOCX已生成: {docx_file}")
